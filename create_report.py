from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_report():
    document = Document()

    # Title
    title = document.add_heading('LLM Evaluation & Model Selection Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    document.add_heading('1. Executive Summary', level=1)
    p = document.add_paragraph()
    p.add_run('The AURA RAG system was evaluated against a comprehensive test suite of 40 diverse queries, covering author lookup, departmental filtering, summarization, and metadata extraction. ').bold = False
    p.add_run('The current system achieved a pass rate of 31/40 (77.5%)').bold = True
    p.add_run(', demonstrating strong capability in semantic search and fact retrieval. This report details the performance of the current local setup and compares it with other viable Large Language Models (LLMs) for future production deployment.')

    # Evaluation Methodology
    document.add_heading('2. Evaluation Methodology', level=1)
    document.add_paragraph('The evaluation dataset ("llm_eval_samples.jsonl") consists of 40 test cases designed to stress-test specific RAG capabilities:')
    
    items = [
        ('Author Lookup', 'Identifying faculty based on research topics (e.g., "Who published on Deep Learning?").'),
        ('Departmental Filtering', 'Isolating publications by department (CSE, ECE, etc.).'),
        ('Semantic Search', 'Finding concepts without exact keyword matches (e.g., "CNN" -> "Deep Learning").'),
        ('Strict Formatting', 'Adhering to output constraints (e.g., "bullet points", "JSON format").'),
        ('Fact Verification', 'Ensuring dates and journal names are accurate.')
    ]
    for key, desc in items:
        p = document.add_paragraph(style='List Bullet')
        p.add_run(f'{key}: ').bold = True
        p.add_run(desc)

    # Current Performance
    document.add_heading('3. Current System Performance', level=1)
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Result'
    hdr_cells[2].text = 'Analysis'
    
    data = [
        ('Pass Rate', '31 / 40 (77.5%)', 'High for a local system. Major improvement from initial baseline (12/40) after enabling semantic embeddings.'),
        ('Semantic Recall', 'Strong', 'Successfully links "CNN" to "Deep Learning" and "Alzheimer" to "MRI" contexts.'),
        ('Formatting', 'Moderate', 'Struggles with strict counting constraints (e.g., "exactly 3 bullets") and complex date filtering.'),
        ('Latency', 'Low (<2s)', 'Local embeddings + lightweight inference provide near real-time responses.')
    ]
    
    for metric, result, analysis in data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = result
        row_cells[2].text = analysis

    # Model Comparison
    document.add_heading('4. LLM Comparison & Recommendations', level=1)
    document.add_paragraph('The following models were analyzed for potential integration:')

    # Comparison Table
    table2 = document.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Model'
    hdr_cells2[1].text = 'Type'
    hdr_cells2[2].text = 'Cost'
    hdr_cells2[3].text = 'Privacy'
    hdr_cells2[4].text = 'Recommended For'

    models = [
        ('Current System (Hybrid)', 'Local + Rule-based', 'Free', 'High (Local)', 'Development / Offline Use'),
        ('Gemini 1.5 Flash', 'Cloud API', 'Low', 'Medium (Cloud)', 'Production (Speed + Context)'),
        ('Llama 3 (8B)', 'Local LLM', 'Free (Hardware)', 'High (Local)', 'Strict Privacy / No Internet'),
        ('GPT-4o', 'Cloud API', 'High', 'Medium (Cloud)', 'Complex Reasoning Tasks')
    ]

    for model, mtype, cost, priv, rec in models:
        row_cells = table2.add_row().cells
        row_cells[0].text = model
        row_cells[1].text = mtype
        row_cells[2].text = cost
        row_cells[3].text = priv
        row_cells[4].text = rec

    # Detailed Analysis of Options
    document.add_heading('4.1 Gemini 1.5 Flash (Recommended for Production)', level=2)
    document.add_paragraph('Google’s Gemini 1.5 Flash is optimized for high-volume, low-latency tasks. Its massive context window (1M tokens) allows it to ingest entire document sets if needed, reducing the complexity of the chunking strategy. It is significantly faster and cheaper than GPT-4o while offering comparable performance for summarization and extraction tasks.')

    document.add_heading('4.2 Llama 3 (8B) - Local Alternative', level=2)
    document.add_paragraph('For a fully offline deployment (e.g., within college intranet without external access), Llama 3 8B is the state-of-the-art open-source model. It requires a GPU with at least 6GB VRAM for decent speed. It outperforms previous 7B models in instruction following, making it suitable for the "strict formatting" tasks where the current system struggles.')

    # Conclusion
    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph('The AURA system is currently performing well (77.5% pass rate) using optimized local embeddings and fallback logic. To bridge the gap to >90% performance—specifically for complex date logic and strict formatting—we recommend integrating **Gemini 1.5 Flash**. This will provide the reasoning capabilities of a large model while maintaining low latency.')

    # Save
    file_path = 'LLM_Evaluation_Report.docx'
    document.save(file_path)
    print(f"Report saved to {file_path}")

if __name__ == "__main__":
    create_report()
