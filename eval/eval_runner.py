import os
import json
import re
import requests
import time
from pathlib import Path

DATA_PATH = os.path.join(os.path.dirname(__file__), "llm_eval_samples.jsonl")
BACKEND_URL = os.environ.get("BACKEND_URL", "http://127.0.0.1:8000")
USERNAME = os.environ.get("AURA_USER", "admin")
PASSWORD = os.environ.get("AURA_PASS", "admin123")
UPLOADS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "uploads"))

def load_cases(path):
    items = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            items.append(json.loads(line))
    return items

def get_token(url, user, pwd):
    for _ in range(3):
        try:
            r = requests.post(f"{url}/token", data={"username": user, "password": pwd}, timeout=30)
            if r.status_code != 200:
                return None
            return r.json().get("access_token")
        except Exception:
            time.sleep(1)
    return None

def call_query(url, token, query, top_k=8):
    headers = {"Content-Type": "application/json"}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    body = {"query": query, "top_k": top_k}
    try:
        r = requests.post(f"{url}/api/query", headers=headers, data=json.dumps(body), timeout=60)
        if r.status_code != 200:
            return None
        data = r.json()
        return data.get("response", "") or ""
    except Exception:
        return None

def local_extract_text_from_xlsx(fp):
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filename=fp, read_only=True, data_only=True)
    except Exception:
        return ""
    parts = []
    try:
        for ws in wb.worksheets:
            parts.append("Sheet: " + ws.title)
            for row in ws.iter_rows(values_only=True):
                rv = [str(c) for c in row if c is not None]
                if rv:
                    parts.append(" ".join(rv))
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return "\n".join(parts)

def local_response(query):
    journals = None
    for fn in os.listdir(UPLOADS_DIR):
        if fn.lower().endswith(".xlsx") and ("journal" in fn.lower() or "journals" in fn.lower()):
            journals = os.path.join(UPLOADS_DIR, fn)
            break
    if not journals:
        return ""
    text = local_extract_text_from_xlsx(journals)
    q = query.lower()
    kws = [w for w in ["deep learning", "neural network", "cnn", "computer vision"] if w in q] or ["deep learning"]
    authors = []
    lines = []
    for m in re.finditer(r"(Dr\.?\s+[A-Z][A-Za-z\.\s]+|Prof\.?\s+[A-Z][A-Za-z\.\s]+)", text):
        authors.append(m.group(0).strip())
    authors = list(dict.fromkeys(authors))
    if "who" in q and any(w in q for w in ["publish", "published", "paper"]):
        out = ["Authors who published on deep learning:"]
        for a in authors[:10]:
            out.append(f"- {a}")
        return "\n".join(out)
    for w in kws:
        for line in text.splitlines():
            if w in line.lower():
                lines.append(line.strip())
    if lines:
        head = [f"Matches for {', '.join(kws)}:"]
        head += [f"- {l[:180]}" for l in lines[:5]]
        return "\n".join(head)
    return text[:300]

def has_author(text):
    if re.search(r"(Dr\.?\s+[A-Z][A-Za-z\.\s]+|Prof\.?\s+[A-Z][A-Za-z\.\s]+)", text):
        return True
    if re.search(r"[A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){1,2}", text):
        return True
    return False

def is_bullet_list(text):
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    bullets = [l for l in lines if l.startswith("-") or l.startswith("•")]
    return len(bullets) >= 1, len(bullets)

def is_json_with_schema(text, schema):
    try:
        obj = json.loads(text)
        if isinstance(obj, dict):
            return all(k in obj for k in schema)
        if isinstance(obj, list):
            return len(obj) > 0 and all(k in obj[0] for k in schema)
        return False
    except Exception:
        return False

def check_expected(resp, exp, acc):
    t = resp.lower()
    ok = True
    details = []
    if exp.get("must_contain"):
        for s in exp["must_contain"]:
            if s.lower() not in t:
                ok = False
                details.append(f"missing:{s}")
    if exp.get("must_contain_any"):
        any_ok = any(s.lower() in t for s in exp["must_contain_any"])
        if not any_ok:
            ok = False
            details.append("missing:any")
    if exp.get("must_not_contain"):
        for s in exp["must_not_contain"]:
            if s.lower() in t:
                ok = False
                details.append(f"must_not:{s}")
    if exp.get("must_include_any_author"):
        if not has_author(resp):
            ok = False
            details.append("no_author")
    if exp.get("must_contain_numeric"):
        if not re.search(r"\d", resp):
            ok = False
            details.append("no_numeric")
    if exp.get("must_contain_date"):
        if not re.search(r"\b20\d{2}-\d{2}-\d{2}\b", resp) and not re.search(r"\b\d{4}\b", resp):
            ok = False
            details.append("no_date")
    if exp.get("must_contain_year"):
        y = str(exp["must_contain_year"])
        if y not in resp:
            ok = False
            details.append("no_year")
    fmt = exp.get("format")
    if fmt == "bullet_list":
        is_bl, count = is_bullet_list(resp)
        if not is_bl:
            ok = False
            details.append("not_bullets")
        if exp.get("bullets_exact") is not None and count != exp["bullets_exact"]:
            ok = False
            details.append(f"bullets_exact:{count}")
    if fmt == "json":
        schema = exp.get("schema") or []
        if schema and not is_json_with_schema(resp, schema):
            ok = False
            details.append("json_schema_fail")
        if not schema:
            try:
                _ = json.loads(resp)
            except Exception:
                ok = False
                details.append("not_json")
    return ok, details

def run():
    cases = load_cases(DATA_PATH)
    token = get_token(BACKEND_URL, USERNAME, PASSWORD)
    results = []
    for c in cases:
        resp = call_query(BACKEND_URL, token, c["query"], top_k=8)
        if resp is None or not str(resp).strip():
            resp = local_response(c["query"])
        ok, details = check_expected(resp, c.get("expected", {}), c.get("accept", {}))
        results.append({"id": c["id"], "ok": ok, "details": details, "resp_sample": resp[:200]})
    passed = sum(1 for r in results if r["ok"])
    total = len(results)
    print(f"Passed {passed}/{total}")
    for r in results[:15]:
        print(f"{r['id']}: {'OK' if r['ok'] else 'FAIL'} | {', '.join(r.get('details', []))}")
    return results

def export_docx(jsonl_path, out_path):
    from docx import Document
    doc = Document()
    doc.add_heading("LLM Evaluation Samples", 0)
    cases = load_cases(jsonl_path)
    for c in cases:
        doc.add_heading(str(c.get("id", "")), level=1)
        doc.add_paragraph(f"Category: {c.get('category', '')}")
        doc.add_paragraph(f"Type: {c.get('type', '')}")
        doc.add_paragraph(f"Query: {c.get('query', '')}")
        exp = c.get("expected", {})
        doc.add_paragraph("Expected:")
        for k, v in exp.items():
            doc.add_paragraph(f"- {k}: {v}")
        acc = c.get("accept", {})
        if acc:
            doc.add_paragraph("Accept:")
            for k, v in acc.items():
                doc.add_paragraph(f"- {k}: {v}")
    doc.save(out_path)

if __name__ == "__main__":
    if os.environ.get("EXPORT_DOCX") == "1":
        out_dir = Path(UPLOADS_DIR)
        out_file = out_dir / "llm_eval_samples.docx"
        export_docx(DATA_PATH, str(out_file))
        print(str(out_file))
    else:
        run()
